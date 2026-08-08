from pathlib import Path
import fitz
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / ".cv-build"
OUT.mkdir(exist_ok=True)

NAVY = RGBColor(8, 35, 58)
BLUE = RGBColor(0, 153, 204)
MUTED = RGBColor(76, 91, 105)
BLACK = RGBColor(20, 25, 30)
LIGHT = "D7E3EA"
FONT = "Arial"
PDF_NAVY = (8 / 255, 35 / 255, 58 / 255)
PDF_BLUE = (0, 153 / 255, 204 / 255)
PDF_MUTED = (76 / 255, 91 / 255, 105 / 255)
PDF_BLACK = (20 / 255, 25 / 255, 30 / 255)


DATA = {
    "es": {
        "file": "CV_mauro_trejo",
        "role": "Senior Full Stack Developer | Mobile, Web & Backend",
        "summary_title": "PERFIL PROFESIONAL",
        "summary": (
            "Desarrollador Full Stack con más de 6 años construyendo productos web, mobile y backend para pagos, "
            "telecomunicaciones, retail y salud. Especializado en React Native, React y .NET/C#, con experiencia en "
            "APIs, microservicios, SQL Server, Azure, AWS y soluciones productivas para Latinoamérica."
        ),
        "experience_title": "EXPERIENCIA",
        "jobs": [
            {
                "role": "Senior Full Stack Developer",
                "company": "PinApp · Real Plaza",
                "date": "Jul 2026 - Actualidad",
                "bullets": [
                    "Desarrollo y evolución de la aplicación mobile de Real Plaza y sus servicios backend con React Native, .NET y C#.",
                    "Integración de APIs REST, SQL Server, Firebase y Azure; participación en arquitectura, decisiones técnicas, CI/CD y mejora continua.",
                ],
            },
            {
                "role": "Full Stack Developer",
                "company": "ASAP Consulting · DIRECTV LATAM",
                "date": "Ago 2025 - Jul 2026",
                "bullets": [
                    "Desarrollo full stack de aplicaciones web, mobile y backend para la operación regional de DIRECTV LATAM.",
                    "Microservicios .NET, APIs e integraciones para órdenes de servicio, técnicos, disponibilidad y configuración operativa; interfaces React + TypeScript y funcionalidades mobile.",
                    "Procesos batch multipaís y herramientas internas para automatizar validaciones y soporte; participación en MiDIRECTV, Optimus y DGO.",
                ],
            },
            {
                "role": "Full Stack Developer",
                "company": "Zoco - Servicios de Pago",
                "date": "Dic 2021 - Jul 2025",
                "bullets": [
                    "Soluciones web, mobile y desktop para pagos, operaciones y eventos; desarrollo y publicación de Zoco, Zoco Tickets y Zoco Scanner.",
                    "Plataforma ASP.NET Core + React con CRM, leads de Meta, asistente con IA, gestión de comercios, liquidaciones, impuestos, compliance, monitoreo, reportes y paneles por rol.",
                    "Herramientas WinForms y .NET 8 para carga masiva y automatización financiera de Excel con reglas, QR, macros y controles de calidad.",
                ],
            },
            {
                "role": "Full Stack Developer",
                "company": "Trinidad Salud",
                "date": "Mar 2021 - Actualidad",
                "bullets": [
                    "Ecosistema de salud con backend y APIs Laravel, panel React, aplicación React Native y sitios institucionales desplegados en AWS.",
                    "Gestión de afiliados, planes, consultas, recetas, facturación, credenciales, reintegros y autorizaciones; autenticación, roles, permisos y PDFs.",
                ],
            },
            {
                "role": "Full Stack Developer independiente",
                "company": "Profesional independiente",
                "date": "Feb 2020 - Actualidad",
                "bullets": [
                    "Diseño y desarrollo integral de soluciones web y mobile: relevamiento, arquitectura, APIs, interfaces, bases de datos, despliegue y mantenimiento.",
                    "Digitalización de procesos y construcción de sistemas de gestión, landing pages y plataformas con React, Laravel, .NET y React Native.",
                ],
            },
        ],
        "skills_title": "TECNOLOGÍAS",
        "skills": [
            ("Backend", ".NET / ASP.NET Core, C#, Laravel, APIs REST, microservicios, Entity Framework"),
            ("Frontend & Mobile", "React, React Native, TypeScript, JavaScript, Expo, HTML, CSS"),
            ("Datos & Cloud", "SQL Server, MySQL, Oracle, Azure, AWS, Firebase, Docker, CI/CD"),
        ],
        "education_title": "EDUCACIÓN",
        "education": "Técnico Universitario en Programación · Universidad Tecnológica Nacional (UTN) · 2018 - 2021",
    },
    "en": {
        "file": "CV_mauro_trejo_en",
        "role": "Senior Full Stack Developer | Mobile, Web & Backend",
        "summary_title": "PROFESSIONAL PROFILE",
        "summary": (
            "Full Stack Developer with 6+ years building web, mobile and backend products for payments, telecommunications, "
            "retail and healthcare. Specialized in React Native, React and .NET/C#, with experience in APIs, microservices, "
            "SQL Server, Azure, AWS and production solutions used across Latin America."
        ),
        "experience_title": "EXPERIENCE",
        "jobs": [
            {
                "role": "Senior Full Stack Developer",
                "company": "PinApp · Real Plaza",
                "date": "Jul 2026 - Present",
                "bullets": [
                    "Development and evolution of the Real Plaza mobile app and its backend services with React Native, .NET and C#.",
                    "REST API, SQL Server, Firebase and Azure integration; contribution to architecture, technical decisions, CI/CD and continuous improvement.",
                ],
            },
            {
                "role": "Full Stack Developer",
                "company": "ASAP Consulting · DIRECTV LATAM",
                "date": "Aug 2025 - Jul 2026",
                "bullets": [
                    "Full stack development of web, mobile and backend applications supporting DIRECTV LATAM's regional operations.",
                    ".NET microservices, APIs and integrations for service orders, technicians, availability and operational configuration; React + TypeScript interfaces and mobile functionality.",
                    "Multi-country batch processes and internal tools for automated validations and support; contributions to MiDIRECTV, Optimus and DGO.",
                ],
            },
            {
                "role": "Full Stack Developer",
                "company": "Zoco - Payment Services",
                "date": "Dec 2021 - Jul 2025",
                "bullets": [
                    "Web, mobile and desktop solutions for payments, operations and events, including the development and release of Zoco, Zoco Tickets and Zoco Scanner.",
                    "ASP.NET Core + React platform with CRM, Meta leads, an AI assistant, merchant management, settlements, taxes, compliance, monitoring, reporting and role-based panels.",
                    "WinForms and .NET 8 tools for bulk uploads and financial Excel automation with rules, QR workflows, macros and quality controls.",
                ],
            },
            {
                "role": "Full Stack Developer",
                "company": "Trinidad Salud",
                "date": "Mar 2021 - Present",
                "bullets": [
                    "Healthcare ecosystem with a Laravel backend and APIs, React admin panel, React Native app and institutional websites deployed on AWS.",
                    "Member and plan management, appointments, prescriptions, billing, digital cards, reimbursements and authorizations; authentication, roles, permissions and PDFs.",
                ],
            },
            {
                "role": "Independent Full Stack Developer",
                "company": "Freelance",
                "date": "Feb 2020 - Present",
                "bullets": [
                    "End-to-end web and mobile solutions: discovery, architecture, APIs, interfaces, databases, deployment and maintenance.",
                    "Process digitization and development of management systems, landing pages and platforms with React, Laravel, .NET and React Native.",
                ],
            },
        ],
        "skills_title": "TECHNOLOGIES",
        "skills": [
            ("Backend", ".NET / ASP.NET Core, C#, Laravel, REST APIs, microservices, Entity Framework"),
            ("Frontend & Mobile", "React, React Native, TypeScript, JavaScript, Expo, HTML, CSS"),
            ("Data & Cloud", "SQL Server, MySQL, Oracle, Azure, AWS, Firebase, Docker, CI/CD"),
        ],
        "education_title": "EDUCATION",
        "education": "University Programming Technician · Universidad Tecnológica Nacional (UTN) · 2018 - 2021",
    },
}


def set_font(run, size=9, bold=False, color=BLACK, italic=False):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color


def add_bottom_border(paragraph, color=LIGHT, size="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    rel = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel)
    run_el = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "007FA8")
    r_pr.append(color)
    run_el.append(r_pr)
    text_el = OxmlElement("w:t")
    text_el.text = text
    run_el.append(text_el)
    hyperlink.append(run_el)
    paragraph._p.append(hyperlink)


def add_section_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(2.5)
    run = p.add_run(text)
    set_font(run, size=10.2, bold=True, color=NAVY)
    add_bottom_border(p)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.first_line_indent = Inches(-0.12)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0.7)
    p.paragraph_format.line_spacing = 1.0
    set_font(p.add_run(text), size=8.15)


def build(language):
    data = DATA[language]
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.42)
    section.bottom_margin = Inches(0.42)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)
    section.header_distance = Inches(0.2)
    section.footer_distance = Inches(0.2)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(9)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_font(title.add_run("MAURO NICOLÁS TREJO"), size=23, bold=True, color=NAVY)

    role = doc.add_paragraph()
    role.paragraph_format.space_before = Pt(0)
    role.paragraph_format.space_after = Pt(3)
    set_font(role.add_run(data["role"]), size=11, bold=True, color=BLUE)

    contact = doc.add_paragraph()
    contact.paragraph_format.space_after = Pt(3)
    set_font(contact.add_run("Tucumán, Argentina  |  +54 381 476-7206  |  maurocrsiete@gmail.com  |  "), size=8.2, color=MUTED)
    add_hyperlink(contact, "LinkedIn", "https://www.linkedin.com/in/mauro-trejo-98968b1ab/")
    set_font(contact.add_run("  |  "), size=8.2, color=MUTED)
    add_hyperlink(contact, "Portfolio", "https://www.maurotrejoportafolio.com/")

    add_section_heading(doc, data["summary_title"])
    summary = doc.add_paragraph()
    summary.paragraph_format.space_after = Pt(1.5)
    summary.paragraph_format.line_spacing = 1.05
    set_font(summary.add_run(data["summary"]), size=8.45)

    add_section_heading(doc, data["experience_title"])
    for job in data["jobs"]:
        p = doc.add_paragraph()
        p.paragraph_format.keep_with_next = True
        p.paragraph_format.space_before = Pt(2.2)
        p.paragraph_format.space_after = Pt(0.5)
        set_font(p.add_run(job["role"]), size=9, bold=True, color=NAVY)
        set_font(p.add_run(f"  |  {job['company']}"), size=8.6, bold=True, color=BLUE)
        set_font(p.add_run(f"  |  {job['date']}"), size=8.1, italic=True, color=MUTED)
        for bullet in job["bullets"]:
            add_bullet(doc, bullet)

    add_section_heading(doc, data["skills_title"])
    for label, value in data["skills"]:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0.6)
        p.paragraph_format.line_spacing = 1.0
        set_font(p.add_run(f"{label}: "), size=8.25, bold=True, color=NAVY)
        set_font(p.add_run(value), size=8.25)

    add_section_heading(doc, data["education_title"])
    education = doc.add_paragraph()
    education.paragraph_format.space_after = Pt(0)
    set_font(education.add_run(data["education"]), size=8.4)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(footer.add_run("Mauro Trejo · Senior Full Stack Developer"), size=7.2, color=MUTED)

    out = OUT / f"{data['file']}.docx"
    doc.save(out)
    return out


def _wrap(text, font, size, width):
    words = text.split()
    lines, current = [], ""
    for word in words:
        candidate = f"{current} {word}".strip()
        if fitz.get_text_length(candidate, fontname=font, fontsize=size) <= width:
            current = candidate
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def build_pdf(language):
    data = DATA[language]
    pdf = fitz.open()
    page = pdf.new_page(width=595, height=842)
    left, right, y = 31, 564, 77
    gray = (0.34, 0.34, 0.34)
    dark = (0.18, 0.18, 0.18)

    def line(text, x, baseline, size, font="helv", color=PDF_BLACK):
        page.insert_text((x, baseline), text, fontname=font, fontsize=size, color=color)

    def section(title, current_y):
        current_y += 8
        line(title, left, current_y, 10.2, "hebo", dark)
        page.draw_line((left, current_y + 6), (right, current_y + 6), color=gray, width=0.85)
        return current_y + 17

    def paragraph(text, current_y, size=7.6, width=None, leading=8.9, x=left, font="helv", color=gray):
        width = width or (right - x)
        for wrapped in _wrap(text, font, size, width):
            line(wrapped, x, current_y, size, font, color)
            current_y += leading
        return current_y

    # Keep the original monochrome visual language: large name, italic role,
    # compact icon/contact rail and thin section rules.
    line("Mauro Nicolas Trejo", left + 1.2, y + 1.2, 27, "hebo", (0.72, 0.72, 0.72))
    line("Mauro Nicolas Trejo", left, y, 27, "hebo", dark)
    y += 18
    subtitle = "Programador Full Stack" if language == "es" else "Full Stack Developer"
    line(subtitle, left, y, 10.5, "heit", gray)
    y += 26

    contacts = [
        ("@", "maurocrsiete@gmail.com", "mailto:maurocrsiete@gmail.com"),
        ("T", "+54 381 4767206", "https://wa.me/543814767206"),
        ("P", "Tucumán, Argentina", None),
        ("in", "Mauro Trejo", "https://www.linkedin.com/in/mauro-trejo-98968b1ab/"),
        ("Pf", "maurotrejoportafolio.com", "https://www.maurotrejoportafolio.com/"),
    ]
    x = left
    for icon, text_value, uri in contacts:
        box_w = 11 if len(icon) == 1 else 14
        page.draw_rect(fitz.Rect(x, y - 9, x + box_w, y + 1), color=gray, fill=gray, width=0.5)
        line(icon, x + 2, y - 1.5, 5.5, "hebo", (1, 1, 1))
        text_x = x + box_w + 6
        line(text_value, text_x, y, 7.5, "heit", gray)
        text_w = fitz.get_text_length(text_value, fontname="heit", fontsize=7.5)
        if uri:
            page.insert_link({"kind": fitz.LINK_URI, "from": fitz.Rect(text_x, y - 9, text_x + text_w, y + 2), "uri": uri})
        x = text_x + text_w + 16

    y = section(data["summary_title"], y + 11)
    y = paragraph(data["summary"], y, size=7.55, leading=8.7)
    y = section(data["experience_title"], y + 3)
    for job in data["jobs"]:
        line(job["role"] + ".", left + 2, y, 9.8, "hebo", gray)
        y += 12
        line(f"{job['company']} | {job['date']}", left + 2, y, 8.6, "heit", gray)
        y += 10
        for bullet in job["bullets"]:
            line("•", left + 5, y, 6.8, "helv", gray)
            y = paragraph(bullet, y, size=7.25, width=right - left - 20, leading=8.1, x=left + 16)
        y += 3

    y = section(data["skills_title"], y - 3)
    for label, value in data["skills"]:
        line(f"{label}:", left + 2, y, 7.4, "hebo", gray)
        value_x = left + 2 + fitz.get_text_length(f"{label}: ", fontname="hebo", fontsize=7.4)
        y = paragraph(value, y, size=7.35, leading=8.25, x=value_x, width=right - value_x)

    y = section(data["education_title"], y)
    y = paragraph(data["education"], y, size=7.55, leading=8.6, x=left + 2)

    soft_title = "HABILIDADES INTERPERSONALES" if language == "es" else "SOFT SKILLS"
    soft = (
        ["Capacidad de análisis y resolución de problemas", "Trabajo colaborativo en equipos técnicos", "Comprensión de necesidades del cliente", "Organización y planificación del desarrollo"]
        if language == "es"
        else ["Analytical thinking and problem-solving", "Teamwork and collaboration in technical environments", "Understanding of client requirements", "Organization and development planning"]
    )
    y = section(soft_title, y + 1)
    col_w = (right - left) / 2
    for index, item in enumerate(soft):
        col = index % 2
        row = index // 2
        item_x = left + 7 + col * col_w
        item_y = y + row * 11
        line("•", item_x, item_y, 6.8, "helv", gray)
        paragraph(item, item_y, size=7.25, leading=8.2, x=item_x + 10, width=col_w - 18)
    y += 22

    if y > 827:
        raise RuntimeError(f"CV content overflow for {language}: y={y}")

    out = OUT / f"{data['file']}.pdf"
    pdf.save(out, garbage=4, deflate=True)
    pdf.close()
    return out


if __name__ == "__main__":
    for lang in DATA:
        print(build(lang))
        print(build_pdf(lang))
